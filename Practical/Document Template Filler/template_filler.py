"""Document Template Filler utility.

Loads JSON/YAML data, injects it into DOCX placeholders, fills PDF form fields,
and optionally regenerates a summary PDF using reportlab.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any, Dict

import yaml
from docx import Document
from PyPDF2 import PdfReader, PdfWriter
from PyPDF2.generic import NameObject
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

PLACEHOLDER_FORMAT = "{{{{{key}}}}}"
MODULE_DIR = Path(__file__).resolve().parent


def load_data(data_path: Path) -> Dict[str, Any]:
    """Load a JSON or YAML file into a dictionary."""
    if not data_path.exists():
        raise FileNotFoundError(f"Data file not found: {data_path}")

    suffix = data_path.suffix.lower()
    with data_path.open("r", encoding="utf-8") as f:
        if suffix == ".json":
            return json.load(f)
        if suffix in {".yml", ".yaml"}:
            return yaml.safe_load(f)
    raise ValueError("Unsupported data format. Use JSON or YAML.")


def _replace_placeholders_in_text(text: str, data: Dict[str, Any]) -> str:
    """Replace placeholders inside a raw text string."""
    updated = text
    for key, value in data.items():
        updated = updated.replace(PLACEHOLDER_FORMAT.format(key=key), str(value))
    return updated


def _replace_placeholders_in_paragraph(paragraph, data: Dict[str, Any]) -> None:
    new_text = _replace_placeholders_in_text(paragraph.text, data)
    if new_text != paragraph.text:
        paragraph.text = new_text


def _replace_placeholders_in_table(table, data: Dict[str, Any]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_placeholders_in_paragraph(paragraph, data)


def create_docx_template(destination: Path) -> None:
    """Generate a sample DOCX template with common placeholders."""
    destination.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("Service Confirmation", 0)
    document.add_paragraph("Dear {{client_name}},")
    document.add_paragraph(
        "Thank you for choosing us for {{service}} on {{appointment_date}} at "
        "{{appointment_time}}."
    )
    document.add_paragraph("Summary: {{summary}}")

    table = document.add_table(rows=5, cols=2)
    rows = [
        ("Service", "{{service}}"),
        ("Date", "{{appointment_date}}"),
        ("Time", "{{appointment_time}}"),
        ("Price", "{{price}}"),
        ("Contact", "{{contact_email}} / {{contact_phone}}"),
    ]
    for idx, (label, placeholder) in enumerate(rows):
        table.rows[idx].cells[0].text = label
        table.rows[idx].cells[1].text = placeholder

    document.add_paragraph("We look forward to serving you.")
    document.add_paragraph("Sincerely,")
    document.add_paragraph("{{company_name}} Team")

    document.save(destination)


def create_pdf_template(destination: Path) -> None:
    """Create a small AcroForm PDF template with named text fields."""
    destination.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(destination), pagesize=LETTER)
    width, height = LETTER
    form = c.acroForm

    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, height - 72, "Service Request Form")

    c.setFont("Helvetica", 11)
    field_specs = [
        ("client_name", "Client Name", 72, height - 110),
        ("service", "Service", 72, height - 150),
        ("appointment_date", "Appointment Date", 72, height - 190),
        ("appointment_time", "Appointment Time", 72, height - 230),
        ("price", "Quoted Price", 72, height - 270),
        ("contact_email", "Contact Email", 72, height - 310),
        ("contact_phone", "Contact Phone", 72, height - 350),
    ]

    for name, label, x, y in field_specs:
        c.drawString(x, y + 16, label)
        form.textfield(
            name=name,
            tooltip=label,
            x=x,
            y=y,
            width=300,
            height=20,
            borderStyle="underlined",
            textColor=None,
            borderWidth=0.5,
        )

    summary_y = height - 410
    c.drawString(72, summary_y + 60, "Summary")
    form.textfield(
        name="summary",
        tooltip="Summary",
        x=72,
        y=summary_y,
        width=440,
        height=50,
        borderStyle="solid",
    )

    c.save()


def generate_sample_templates(templates_dir: Path) -> None:
    """Generate example DOCX and PDF templates if they are missing."""
    docx_path = templates_dir / "letter_template.docx"
    pdf_path = templates_dir / "service_form_template.pdf"

    if not docx_path.exists():
        create_docx_template(docx_path)
        print(f"Created sample DOCX template at {docx_path}")
    else:
        print(f"DOCX template already exists at {docx_path}, leaving unchanged.")

    if not pdf_path.exists():
        create_pdf_template(pdf_path)
        print(f"Created sample PDF template at {pdf_path}")
    else:
        print(f"PDF template already exists at {pdf_path}, leaving unchanged.")


def fill_docx_template(
    template_path: Path, data: Dict[str, Any], output_path: Path
) -> None:
    """Fill placeholders inside a DOCX template and save the result."""
    if not template_path.exists():
        raise FileNotFoundError(f"DOCX template not found: {template_path}")

    document = Document(template_path)

    for paragraph in document.paragraphs:
        _replace_placeholders_in_paragraph(paragraph, data)

    for table in document.tables:
        _replace_placeholders_in_table(table, data)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def fill_pdf_form(template_path: Path, data: Dict[str, Any], output_path: Path) -> None:
    """Fill a form-enabled PDF using PyPDF2."""
    if not template_path.exists():
        raise FileNotFoundError(f"PDF template not found: {template_path}")

    reader = PdfReader(str(template_path))
    if not reader.pages:
        raise ValueError("PDF template contains no pages.")

    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)

    # Preserve AcroForm dictionary so fields remain editable after saving.
    root = reader.trailer.get("/Root")
    root_obj = root.get_object() if root and hasattr(root, "get_object") else root
    if root_obj and "/AcroForm" in root_obj:
        writer._root_object.update({NameObject("/AcroForm"): root_obj["/AcroForm"]})

    field_values = {k: str(v) for k, v in data.items()}
    if reader.get_fields():
        for idx in range(len(writer.pages)):
            writer.update_page_form_field_values(writer.pages[idx], field_values)
    else:
        print(
            "Warning: PDF template has no form fields; writing original pages unchanged."
        )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("wb") as f:
        writer.write(f)


def generate_pdf_summary(data: Dict[str, Any], output_path: Path) -> None:
    """Generate a simple summary PDF report using reportlab."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(output_path), pagesize=LETTER)
    width, height = LETTER

    c.setFont("Helvetica-Bold", 18)
    c.drawString(72, height - 72, "Service Summary")

    c.setFont("Helvetica", 12)
    y = height - 110
    for key in [
        "client_name",
        "service",
        "appointment_date",
        "appointment_time",
        "price",
        "contact_email",
        "contact_phone",
    ]:
        if key not in data:
            continue
        line = f"{key.replace('_', ' ').title()}: {data[key]}"
        c.drawString(72, y, line)
        y -= 18

    c.drawString(72, y - 10, "Summary:")
    text_object = c.beginText(72, y - 30)
    text_object.textLines(str(data.get("summary", "")))
    c.drawText(text_object)

    c.save()


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Fill DOCX/PDF templates using JSON or YAML data."
    )
    parser.add_argument(
        "--data",
        required=True,
        type=Path,
        help="Path to JSON/YAML file containing fields.",
    )
    parser.add_argument(
        "--docx-template", type=Path, help="DOCX template path with {{placeholders}}."
    )
    parser.add_argument(
        "--docx-output", type=Path, help="Destination path for filled DOCX."
    )
    parser.add_argument(
        "--pdf-template", type=Path, help="Form-enabled PDF template path."
    )
    parser.add_argument(
        "--pdf-output", type=Path, help="Destination path for filled PDF form."
    )
    parser.add_argument(
        "--generated-pdf",
        type=Path,
        help="Optional path to generate a summary PDF via reportlab.",
    )
    parser.add_argument(
        "--create-sample-templates",
        action="store_true",
        help="Create sample DOCX/PDF templates inside the templates directory if missing.",
    )
    parser.add_argument(
        "--templates-dir",
        type=Path,
        default=MODULE_DIR / "templates",
        help="Directory where sample templates should be created or located.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()

    if args.create_sample_templates:
        generate_sample_templates(args.templates_dir)
        default_docx = args.templates_dir / "letter_template.docx"
        default_pdf = args.templates_dir / "service_form_template.pdf"
        if args.docx_template is None:
            args.docx_template = default_docx
        if args.pdf_template is None:
            args.pdf_template = default_pdf

    data = load_data(args.data)

    if args.docx_template and args.docx_output:
        fill_docx_template(args.docx_template, data, args.docx_output)
        print(f"DOCX saved to {args.docx_output}")
    elif args.docx_template or args.docx_output:
        raise SystemExit(
            "Provide both --docx-template and --docx-output to generate DOCX output."
        )

    if args.pdf_template and args.pdf_output:
        fill_pdf_form(args.pdf_template, data, args.pdf_output)
        print(f"Filled PDF saved to {args.pdf_output}")
    elif args.pdf_template or args.pdf_output:
        raise SystemExit(
            "Provide both --pdf-template and --pdf-output to generate a filled PDF form."
        )

    if args.generated_pdf:
        generate_pdf_summary(data, args.generated_pdf)
        print(f"Generated summary PDF saved to {args.generated_pdf}")

    if not any([args.docx_template, args.pdf_template, args.generated_pdf]):
        print(
            "No outputs requested. Use --docx-template/--pdf-template/--generated-pdf to produce files."
        )


if __name__ == "__main__":
    main()
